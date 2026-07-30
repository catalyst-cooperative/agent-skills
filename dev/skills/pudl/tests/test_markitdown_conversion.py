"""Tests for the markitdown-based file-to-text conversion pattern documented in
references/data-sources.md, under "Blank forms and filer instructions" -> "Reading
them".

markitdown can't currently join this repository's shared pixi dev environment: its
transitive dependency chain (magika -> onnxruntime) has no wheel compatible with
pixi's default osx-arm64 platform-tag baseline for Python 3.14, even though a plain
`uv pip install "markitdown[pdf]"` on this same machine resolves and runs it fine.
That gap is exactly why the skill documents installing markitdown on demand in a
throwaway environment rather than assuming it's preinstalled -- so these tests build
that same throwaway environment with `uv` (already a pixi dependency) and exercise the
documented command, proving the pattern works rather than assuming it does.

Two tests make a real network call, downloading the smallest real PDF and HTML "blank
form"/"instructions" documents linked from live PUDL source documentation pages, to
prove the conversion pattern against real PUDL resources (see references/data-sources.md
for how these URLs are normally discovered). The other two tests convert locally
generated .xlsx and .docx files -- binary, non-text formats PUDL doesn't itself
distribute as documentation -- with no network call, to cover formats beyond PDF/HTML
that agents may still run into (e.g. a user-supplied Word memo or spreadsheet).

Run:  pixi run pytest dev/skills/pudl/tests/test_markitdown_conversion.py -v
"""

import subprocess
import urllib.request
from pathlib import Path

import pytest

# Smallest instructions PDF found under any source's "Download additional
# documentation" section (~120 KB) -- keeps the live download and PDF conversion fast.
FERC714_INSTRUCTIONS_PDF_URL = (
    "https://docs.catalyst.coop/pudl/en/nightly/_downloads/"
    "f48d006b011903ae36c2e6a5acfb4f81/ferc714_instructions_2021-04-16.pdf"
)

# Smallest of the newest-edition (2025-07-31) HTML blank forms linked from ferc1.html's
# "Download additional documentation" section (~1.3 MB).
FERC3Q_ELECTRIC_HTML_URL = (
    "https://docs.catalyst.coop/pudl/en/nightly/_downloads/"
    "5fe9aad5d45a9f43e5b7a65cf0ad5e97/ferc3q_electric_2025-07-31.html"
)


@pytest.fixture(scope="module")
def markitdown_venv(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Build the throwaway venv the documented pattern calls for, once per module.

    references/data-sources.md documents `uv pip install "markitdown[pdf,docx]"` into
    a scratch environment rather than assuming markitdown is already present. This
    fixture runs that pattern (plus the [xlsx] extra and python-docx, needed only to
    generate the .xlsx/.docx fixtures below, not to run markitdown itself).
    """
    venv_dir = tmp_path_factory.mktemp("markitdown-venv")
    subprocess.run(["uv", "venv", str(venv_dir)], check=True, capture_output=True)
    subprocess.run(
        [
            "uv",
            "pip",
            "install",
            "--python",
            str(venv_dir),
            "markitdown[pdf,xlsx,docx]",
            "python-docx",
        ],
        check=True,
        capture_output=True,
    )
    return venv_dir


def _convert(markitdown_venv: Path, path: Path) -> str:
    result = subprocess.run(
        [str(markitdown_venv / "bin" / "markitdown"), str(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout


def _download(url: str, destination: Path) -> None:
    # docs.catalyst.coop returns 403 for urllib's default User-Agent string; a
    # non-empty one (as curl and browsers send by default) is accepted.
    request = urllib.request.Request(url, headers={"User-Agent": "pudl-skill-tests"})
    with urllib.request.urlopen(request, timeout=30) as response:
        destination.write_bytes(response.read())


def test_markitdown_converts_pdf(markitdown_venv: Path, tmp_path: Path) -> None:
    """A real FERC filer-instructions PDF converts to readable markdown text."""
    pdf_path = tmp_path / "ferc714_instructions.pdf"
    _download(FERC714_INSTRUCTIONS_PDF_URL, pdf_path)

    text = _convert(markitdown_venv, pdf_path)

    assert "Form 714" in text
    assert "Instructions" in text


def test_markitdown_converts_html(markitdown_venv: Path, tmp_path: Path) -> None:
    """A real FERC blank-form HTML edition converts to readable markdown text."""
    html_path = tmp_path / "ferc3q_electric.html"
    _download(FERC3Q_ELECTRIC_HTML_URL, html_path)

    text = _convert(markitdown_venv, html_path)

    assert "FERC FORM No. 1" in text


def test_markitdown_converts_xlsx(markitdown_venv: Path, tmp_path: Path) -> None:
    """A locally generated .xlsx -- a binary, non-text format -- also converts
    cleanly, showing the same tool handles more than just PDF and HTML."""
    xlsx_path = tmp_path / "sample.xlsx"
    build_script = (
        "import openpyxl\n"
        "wb = openpyxl.Workbook()\n"
        "ws = wb.active\n"
        "ws['A1'] = 'plant'\n"
        "ws['B1'] = 'capacity_mw'\n"
        "ws['A2'] = 'Comanche'\n"
        "ws['B2'] = 100.5\n"
        f"wb.save({str(xlsx_path)!r})\n"
    )
    subprocess.run(
        [str(markitdown_venv / "bin" / "python"), "-c", build_script],
        check=True,
        capture_output=True,
    )

    text = _convert(markitdown_venv, xlsx_path)

    assert "Comanche" in text
    assert "100.5" in text


def test_markitdown_converts_docx(markitdown_venv: Path, tmp_path: Path) -> None:
    """A locally generated .docx -- another binary, non-text format -- also converts
    cleanly, e.g. for a user-supplied Word memo or errata document."""
    docx_path = tmp_path / "sample.docx"
    build_script = (
        "import docx\n"
        "d = docx.Document()\n"
        "d.add_heading('FERC Form 1 Errata', level=1)\n"
        "d.add_paragraph('Schedule 204 line 5 should reference Comanche station.')\n"
        f"d.save({str(docx_path)!r})\n"
    )
    subprocess.run(
        [str(markitdown_venv / "bin" / "python"), "-c", build_script],
        check=True,
        capture_output=True,
    )

    text = _convert(markitdown_venv, docx_path)

    assert "FERC Form 1 Errata" in text
    assert "Comanche" in text
